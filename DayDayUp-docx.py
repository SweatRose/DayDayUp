# -*- coding = utf-8 -*-
# @Time : 2022-10-13 19:09:29
# @updateTime: 2023-12-20 09:42:52
# @Author : Anonymous
# @File : DayDayUp.py
# @Software : PyCharm

from nntplib import ArticleInfo
import requests
from bs4 import BeautifulSoup                       # 网页解析，获取数据
import time                                         # 时间处理
import re                                           # 正则匹配
import xlwt                                         # 进行excel操作
from pyquery import PyQuery as pq
import os
import shutil
#操作word的库
from docx import Document
from io import BytesIO
from w3lib.html import remove_tags
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 设置对象居中、对齐等。
from docx.shared import Inches  # 设置图像大小
from docx.shared import Pt  # 设置像素、缩进等
from docx.shared import RGBColor  # 设置字体颜色

def main():
    datalist = []                                   # 保存文章标题和内容
    theme_list = {'4006': '申论热点', '4005': '申论范文', '4007': '申论技巧', '3994': '全部申论',
                  '3993': '行测技巧', '3995': '面试技巧', '4011': '面试热点', '4012': '公安基础知识', '3997': '综合指导'}
    theme = input("请选择爬取内容代码（申论热点：4006、申论范文：4005、申论技巧：4007、全部申论：3994、行测技巧：3993、面试技巧：3995、面试热点：4011、公安基础知识：4012、综合指导：3997）：")
    page = input("请输入需要爬取的页数：")
    URL_list = get_pages_url(theme, int(page))              # 获取每个文章的URL
    for i in range(len(URL_list)):
        print("正在获取第{}篇范文".format(i+1))
        # datalist.extend(get_Data(URL_list[i]))
        title, content = get_Data(URL_list[i])  # 获取文章标题和内容
        datalist.append((title, content))  # 将标题和内容作为元组加入列表
        time.sleep(0.5)
        i += 1

        saveData(datalist, theme_list[theme])
        # saveData(title, content)


def askURL(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/63.0.3239.132 Safari/537.36'
    }

    response = requests.get(url=url, headers=headers)
    html = response.content.decode("utf-8")
    bs = BeautifulSoup(html, "html.parser")  # 解析html，提取数据
    return bs

# 获取每一页的全部url
def get_pages_url(theme, page):
    data = []
    for num in range(0, page, 1):
        url = 'http://www.offcn.com/gwy/ziliao/{}/{}.html'.format(theme, num+1)
        print("正在提取第" + str(num+1) + "页")
        bs = askURL(url)
        time.sleep(0.5)
        list01 = str(bs.find_all('ul', class_="lh_newBobotm02"))
        pat_link = r'<a href="(.*?)" target="_blank" title=.*?'
        link = re.findall(pat_link, list01, re.S)
        link_list = str(link).strip("[").strip("]").replace("'", "").split(",")
        data.extend(link_list)
    return data

def get_Data(url):
    articleData = []
    bs = askURL(url)
    title_MsgStr = bs.select("h1")
    title = ""
    for title_str in title_MsgStr:
        title = title + title_str.text
    articleData.append(title)

    content = ""
    list01 = str(bs.find_all('div', class_="offcn_shocont"))
    content = re.sub(r'<.*?>', '', list01)  # 去除HTML标签
    content = re.sub(r'\[.*?\]', '', content)  # 去除方括号标签
    articleData.append(str(content))
    return articleData

def clean_filename(filename):
    cleaned_filename = re.sub(r'[\\/*?:"<>|]', '', filename)  # 去除文件名中的非法字符
    cleaned_filename = cleaned_filename.replace('(进入阅读模式)', '')  # 去除标题中的"(进入阅读模式)"
    return cleaned_filename.strip()  # 去除首尾空格

def saveData(datalist, theme_name):
    for i, data in enumerate(datalist):
        doc = Document()
        title = data[0].replace('(进入阅读模式)', '')  # 获取标题并去除"(进入阅读模式)"
        title = re.sub('进入阅读模式', '', title)  # 删除标题中的特殊字符
        title = re.sub('([^\u4e00-\u9fa5\d])', '', title)  # 删除标题中的特殊字符
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.name = '微软雅黑'
        title_run.font.size = Pt(24)  # 设置标题大小
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # 字体颜色
        title_paragraph.alignment = 1  # 设置标题居中
        


        content = data[1].replace('[', '').replace(']', '')  # 去除正文中的方括号及其中的内容
        content = re.sub('进入阅读模式', '', content)  # 删除标题中的特殊字符
        # content = re.sub(r'&[gl]t;.*?\|.*?元.*?核心考点', '', content)
        # content = re.sub(r'&gt;.*?$', '', content, flags=re.MULTILINE)
        content_paragraph = doc.add_paragraph()
        content_run = content_paragraph.add_run(content)
        content_run.font.name = '仿宋'
        content_run.font.size = Pt(16)

        # 添加图片
        try:
            image_urls = data[2]  # 假设data[2]是图片的URL列表
            for img_url in image_urls:
                response = requests.get(img_url)
                image_stream = BytesIO(response.content)
                doc.add_picture(image_stream, width=Inches(4))  # 设置图片宽度为4英寸
        except IndexError:
            pass  # 如果没有图片URL列表，跳过添加图片的操作
        

        doc.save('./archive/' + theme_name + '-' + clean_filename(title)  +'.docx')  # 使用去除"(进入阅读模式)"后的标题作为文件名保存


if __name__ == '__main__':
    folder = './archive'
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
    main()
    print("爬取完毕！！！")